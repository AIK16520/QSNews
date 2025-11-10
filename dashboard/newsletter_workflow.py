"""
Simplified Newsletter Workflow
Direct preview with AI-generated link descriptions and export options
"""

import streamlit as st
import openai
import os
from datetime import datetime
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_link_description(link_data: dict, newsletter_context: str = "") -> str:
    """
    Generate a 2-sentence description for a link using AI.

    Args:
        link_data: Dictionary with 'title', 'url', 'context', 'explanation'
        newsletter_context: Optional context from the newsletter

    Returns:
        2-sentence description
    """
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        # Fallback without AI
        title = link_data.get('title', 'Link')
        context = link_data.get('context', '') or link_data.get('explanation', '')
        if context:
            return f"{title}. {context[:150]}"
        return f"{title}. More details available at the source."

    try:
        client = openai.OpenAI(api_key=api_key)

        title = link_data.get('title', 'Link')
        url = link_data.get('url', '')
        context = link_data.get('context', '') or link_data.get('explanation', '')

        prompt = f"""Write exactly 2 concise sentences (max 40 words total) about this link from an AI/tech newsletter:

Title: {title}
Context: {context}
{f"Newsletter context: {newsletter_context[:200]}" if newsletter_context else ""}

Write 2 clear, informative sentences that explain what this link is about and why it's relevant to AI/tech professionals. Do not include the URL."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a tech newsletter editor. Write concise, informative 2-sentence descriptions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=100
        )

        description = response.choices[0].message.content.strip()
        return description

    except Exception as e:
        # Fallback on error
        title = link_data.get('title', 'Link')
        context = link_data.get('context', '') or link_data.get('explanation', '')
        if context:
            return f"{title}. {context[:150]}"
        return f"{title}. More details available at the source."


def export_as_pdf(content: str, selected_links: list, filename: str = None) -> BytesIO:
    """Export content as PDF"""
    buffer = BytesIO()

    if not filename:
        filename = f"newsletter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#1a1f2e',
        spaceAfter=12
    )

    # Body style
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=10
    )

    # Link title style
    link_title_style = ParagraphStyle(
        'LinkTitle',
        parent=styles['Heading2'],
        fontSize=13,
        textColor='#6bb6ff',
        spaceAfter=6
    )

    # Add title
    story.append(Paragraph("Newsletter Summary", title_style))
    story.append(Spacer(1, 0.2*inch))

    # Add selected links
    for i, link in enumerate(selected_links, 1):
        # Link title
        title_text = f"{i}. {link['title']}"
        story.append(Paragraph(title_text, link_title_style))

        # Description
        description = link.get('description', '')
        story.append(Paragraph(description, body_style))

        # URL
        url_text = f'<link href="{link["url"]}">{link["url"]}</link>'
        story.append(Paragraph(url_text, body_style))
        story.append(Spacer(1, 0.15*inch))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


def export_as_docx(content: str, selected_links: list, filename: str = None) -> BytesIO:
    """Export content as DOCX"""
    buffer = BytesIO()

    if not filename:
        filename = f"newsletter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc = Document()

    # Add title
    title = doc.add_heading('Newsletter Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Add selected links
    for i, link in enumerate(selected_links, 1):
        # Link title (as heading)
        heading = doc.add_heading(f"{i}. {link['title']}", level=2)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(107, 182, 255)

        # Description
        description = link.get('description', '')
        desc_para = doc.add_paragraph(description)
        desc_para.paragraph_format.space_after = Pt(6)

        # URL (as hyperlink)
        url_para = doc.add_paragraph()
        url_run = url_para.add_run(link['url'])
        url_run.font.size = Pt(10)
        url_run.font.color.rgb = RGBColor(107, 182, 255)
        url_para.paragraph_format.space_after = Pt(12)

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def render_newsletter_preview_page(session, newsletter_id: int):
    """
    Simplified preview page with AI-generated descriptions and checkboxes.

    Args:
        session: Database session
        newsletter_id: ID of the newsletter to preview
    """
    from src.utils.database import Newsletter

    # Get the newsletter
    newsletter = session.query(Newsletter).filter_by(id=newsletter_id).first()

    if not newsletter:
        st.error("Newsletter not found")
        return

    st.title(f"Preview: {newsletter.title}")

    # Show newsletter info
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown(f"**Source:** {newsletter.source}")
        st.markdown(f"**Date:** {newsletter.published_date.strftime('%Y-%m-%d') if newsletter.published_date else 'N/A'}")
    with col2:
        if st.button("← Back to Review", type="secondary"):
            newsletter.status = 'not_included'
            session.commit()
            st.session_state.current_page = 'review'
            st.rerun()

    st.markdown("---")

    # Get extracted links
    extracted_links = newsletter.extracted_links if newsletter.extracted_links else []

    if not extracted_links:
        st.warning("No links found in this newsletter")
        return

    st.subheader(f"Links ({len(extracted_links)})")
    st.markdown("Select links to include in your newsletter summary:")

    # Initialize session state for selections and descriptions
    if f'selected_links_{newsletter_id}' not in st.session_state:
        st.session_state[f'selected_links_{newsletter_id}'] = []

    if f'link_descriptions_{newsletter_id}' not in st.session_state:
        st.session_state[f'link_descriptions_{newsletter_id}'] = {}

    # Generate descriptions button
    col1, col2 = st.columns([3, 1])
    with col2:
        if st.button("🤖 Generate All Descriptions", type="primary", use_container_width=True):
            with st.spinner("Generating AI descriptions for all links..."):
                newsletter_context = newsletter.summary or ""
                for link in extracted_links:
                    link_key = link.get('url', '')
                    if link_key and link_key not in st.session_state[f'link_descriptions_{newsletter_id}']:
                        description = generate_link_description(link, newsletter_context)
                        st.session_state[f'link_descriptions_{newsletter_id}'][link_key] = description
            st.success("✅ Descriptions generated!")
            st.rerun()

    # Display links with checkboxes
    for i, link in enumerate(extracted_links):
        link_url = link.get('url', '')
        link_title = link.get('title', f'Link {i+1}')

        # Checkbox for selection
        is_selected = link_url in st.session_state[f'selected_links_{newsletter_id}']

        col1, col2 = st.columns([0.05, 0.95])

        with col1:
            selected = st.checkbox(
                "Select",
                value=is_selected,
                key=f"checkbox_{newsletter_id}_{i}",
                label_visibility="collapsed"
            )

            # Update selection
            if selected and link_url not in st.session_state[f'selected_links_{newsletter_id}']:
                st.session_state[f'selected_links_{newsletter_id}'].append(link_url)
            elif not selected and link_url in st.session_state[f'selected_links_{newsletter_id}']:
                st.session_state[f'selected_links_{newsletter_id}'].remove(link_url)

        with col2:
            with st.expander(f"**{i+1}. {link_title}**", expanded=False):
                st.markdown(f"**URL:** [{link_url}]({link_url})")

                # Show description or generate button
                description = st.session_state[f'link_descriptions_{newsletter_id}'].get(link_url, '')

                if description:
                    st.markdown(f"**Description:**")
                    st.info(description)
                else:
                    if st.button(f"Generate Description", key=f"gen_{newsletter_id}_{i}"):
                        with st.spinner("Generating..."):
                            newsletter_context = newsletter.summary or ""
                            description = generate_link_description(link, newsletter_context)
                            st.session_state[f'link_descriptions_{newsletter_id}'][link_url] = description
                        st.rerun()

    st.markdown("---")

    # Build preview of selected links
    st.subheader(f"Preview ({len(st.session_state[f'selected_links_{newsletter_id}'])} links selected)")

    selected_link_objects = []
    preview_text = ""

    for link_url in st.session_state[f'selected_links_{newsletter_id}']:
        # Find the link object
        link_obj = next((l for l in extracted_links if l.get('url') == link_url), None)
        if link_obj:
            description = st.session_state[f'link_descriptions_{newsletter_id}'].get(link_url, '')
            link_obj['description'] = description
            selected_link_objects.append(link_obj)

            # Add to preview text
            preview_text += f"**{link_obj.get('title', 'Link')}**\n\n"
            preview_text += f"{description}\n\n"
            preview_text += f"{link_url}\n\n"
            preview_text += "---\n\n"

    # Read-only preview
    if preview_text:
        st.markdown(preview_text)
    else:
        st.info("No links selected yet. Check the boxes above to select links for your newsletter.")

    # Export buttons
    if selected_link_objects:
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 1, 2])

        with col1:
            # Export as PDF
            pdf_buffer = export_as_pdf(preview_text, selected_link_objects)
            st.download_button(
                label="📄 Export as PDF",
                data=pdf_buffer,
                file_name=f"newsletter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
                type="primary",
                use_container_width=True
            )

        with col2:
            # Export as DOCX
            docx_buffer = export_as_docx(preview_text, selected_link_objects)
            st.download_button(
                label="📝 Export as DOC",
                data=docx_buffer,
                file_name=f"newsletter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )

# -*- coding: utf-8 -*-
"""
Streamlit Dashboard for AI Newsletter Article Review
Provides an interactive UI for reviewing and managing articles.
"""

import sys
import os

# Add parent directory to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import streamlit as st
from streamlit.components.v1 import html as st_html
import pandas as pd
from datetime import datetime, timedelta
import html
import re

from src.utils.database import get_session, Article, Category, Industry, Newsletter
from src.processors.report_generator import generate_pending_review, generate_newsletter_draft
from src.processors.article_generator import generate_standard_format, generate_article_with_ai
from src.fetchers.rss_fetcher import fetch_all_rss, save_fetched_articles
from src.processors.pipeline import process_articles
from src.processors.newsletter_pipeline import fetch_and_process_gmail_newsletters
import openai

# Import simplified newsletter workflow
from dashboard.newsletter_workflow import render_newsletter_preview_page
from dashboard.builder_workflow import render_newsletter_builder_workflow
from dashboard.newsletter_formatter import generate_strictlyvc_style_newsletter

# Page config
st.set_page_config(
    page_title="AI Newsletter - Article Review",
    page_icon=":newspaper:",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'refresh_trigger' not in st.session_state:
    st.session_state.refresh_trigger = 0
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'review'
if 'applied_filters' not in st.session_state:
    st.session_state.applied_filters = None
if 'filter_applied' not in st.session_state:
    st.session_state.filter_applied = False
if 'articles_to_show' not in st.session_state:
    st.session_state.articles_to_show = 20  # Show 20 articles initially
if 'newsletters_to_show' not in st.session_state:
    st.session_state.newsletters_to_show = 50  # Show 50 newsletters initially
if 'builder_items_to_show' not in st.session_state:
    st.session_state.builder_items_to_show = 50  # Show 50 items in builder initially
if 'content_type' not in st.session_state:
    st.session_state.content_type = 'newsletters'  # 'newsletter_builder', 'articles', or 'newsletters'
if 'preview_newsletter_id' not in st.session_state:
    st.session_state.preview_newsletter_id = None

# Custom CSS - Modern Dark Theme
st.markdown("""
<style>
    /* Global theme */
    .main {
        background-color: #0f1419 !important;
    }

    [data-testid="collapsedControl"] {
        display: none;
    }

    section[data-testid="stSidebar"] {
        width: 300px !important;
        background-color: #1a1f2e !important;
    }

    /* Sidebar styling */
    section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] {
        color: #e4e6eb !important;
    }

    /* Newsletter Card Styling */
    .newsletter-card {
        background: linear-gradient(135deg, #1a1f2e 0%, #141824 100%);
        border: 1px solid rgba(107, 182, 255, 0.15);
        border-radius: 16px;
        padding: 28px;
        margin-bottom: 24px;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
    }

    .newsletter-card:hover {
        border-color: rgba(107, 182, 255, 0.35);
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.5);
        transform: translateY(-2px);
    }

    .newsletter-header-row {
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        margin-bottom: 16px;
    }

    .newsletter-title-link {
        color: #6bb6ff !important;
        font-size: 22px;
        font-weight: 600;
        text-decoration: none;
        line-height: 1.4;
        transition: color 0.2s;
    }

    .newsletter-title-link:hover {
        color: #8cc9ff !important;
        text-decoration: underline;
    }

    .newsletter-meta-info {
        color: #8b92a8;
        font-size: 14px;
        margin: 8px 0 16px 0;
        display: flex;
        align-items: center;
        gap: 12px;
    }

    .newsletter-source {
        font-weight: 500;
    }

    .newsletter-date {
        display: flex;
        align-items: center;
        gap: 6px;
    }

    .status-badge {
        background-color: #2d4a3e;
        color: #6db380;
        padding: 6px 14px;
        border-radius: 14px;
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }

    .status-badge-read {
        background-color: #2d3b4a;
        color: #7ba3c7;
    }

    .newsletter-preview-text {
        color: #e4e6eb;
        font-size: 15px;
        line-height: 1.7;
        margin-bottom: 16px;
    }

    .newsletter-preview-text a {
        color: #6bb6ff !important;
        text-decoration: none;
        border-bottom: 1px solid rgba(107, 182, 255, 0.3);
        transition: all 0.2s;
    }

    .newsletter-preview-text a:hover {
        color: #8cc9ff !important;
        border-bottom-color: rgba(140, 201, 255, 0.6);
    }

    .tag-container {
        display: flex;
        flex-wrap: wrap;
        gap: 8px;
        margin: 0;
    }

    .tag {
        background-color: #252b3b;
        color: #8b92a8;
        padding: 6px 12px;
        border-radius: 6px;
        font-size: 12px;
        font-weight: 500;
    }

    .tag.tools {
        background-color: #2d4a3e;
        color: #6db380;
    }

    .tag.legal {
        background-color: #4a2d3e;
        color: #c77ba3;
    }

    .tag.funding {
        background-color: #2d3b4a;
        color: #7ba3c7;
    }

    .tag.security {
        background-color: #4a3e2d;
        color: #c7a37b;
    }

    .tag.events {
        background-color: #3e4a2d;
        color: #a3c77b;
    }

    .tag.research {
        background-color: #3e2d4a;
        color: #b37bc7;
    }

    /* Custom button classes removed - using Streamlit defaults */

    /* Newsletter toggle button colors removed - using Streamlit defaults */

    /* Expander styling for newsletter cards */
    .streamlit-expanderHeader {
        background-color: rgba(255, 255, 255, 0.05) !important;
        border-radius: 8px !important;
        color: #8b92a8 !important;
        font-size: 14px !important;
        padding: 12px 16px !important;
        margin-top: 16px !important;
    }

    .streamlit-expanderHeader:hover {
        background-color: rgba(255, 255, 255, 0.08) !important;
    }

    /* Style for status indicators in expander titles */
    .status-included {
        color: #6db380 !important;
        font-weight: 500;
        background-color: #2d4a3e;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .status-not-included {
        color: #8b92a8 !important;
        font-weight: 500;
        background-color: #252b3b;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    /* Button styling removed - using Streamlit defaults */

    /* Newsletter container styling - target all containers in newsletter section */
    div[data-testid="stVerticalBlock"] > div[data-testid="element-container"]:has(div[data-testid="stMarkdownContainer"] h3) {
        background-color: rgba(255, 255, 255, 0.03);
        border-radius: 12px;
        padding: 24px;
        margin-bottom: 24px;
        border: 1px solid rgba(255, 255, 255, 0.15);
        transition: all 0.3s ease;
    }

    div[data-testid="stVerticalBlock"] > div[data-testid="element-container"]:has(div[data-testid="stMarkdownContainer"] h3):hover {
        background-color: rgba(255, 255, 255, 0.05);
        border-color: rgba(255, 255, 255, 0.25);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
    }

    /* Newsletter title styling */
    .newsletter-title {
        font-size: 1.3em !important;
        font-weight: 600 !important;
        margin-bottom: 8px !important;
        line-height: 1.4 !important;
    }
    .newsletter-title a {
        color: #6bb6ff !important;
        text-decoration: none !important;
    }
    .newsletter-title a:hover {
        color: #8cc9ff !important;
        text-decoration: underline !important;
    }

    /* Newsletter meta (source & date) */
    .newsletter-meta {
        color: rgba(255, 255, 255, 0.6) !important;
        font-size: 0.9em !important;
        margin-bottom: 12px !important;
    }

    /* Newsletter summary */
    .newsletter-summary {
        font-size: 1.05em !important;
        line-height: 1.7 !important;
        color: rgba(255, 255, 255, 0.9) !important;
    }
    .newsletter-summary a {
        color: #6bb6ff !important;
        text-decoration: none !important;
        border-bottom: 1px solid rgba(107, 182, 255, 0.3);
    }
    .newsletter-summary a:hover {
        color: #8cc9ff !important;
        border-bottom-color: rgba(140, 201, 255, 0.6);
    }

    /* Newsletter section general text color */
    div[data-testid="stMarkdownContainer"] p {
        color: rgba(255, 255, 255, 0.9) !important;
    }

    /* Expander styling */
    .streamlit-expanderHeader {
        background-color: rgba(255, 255, 255, 0.05) !important;
        border-radius: 8px !important;
        padding: 12px 16px !important;
        margin-top: 16px !important;
        transition: all 0.2s ease !important;
    }
    .streamlit-expanderHeader:hover {
        background-color: rgba(255, 255, 255, 0.08) !important;
    }

    /* Header count styling */
    .newsletter-count {
        color: rgba(255, 255, 255, 0.5);
        font-size: 0.9em;
        font-weight: normal;
    }

    /* Toggle switch styling - green for included, red for not included */
    /* Override Streamlit's default toggle colors */
    label[data-baseweb="checkbox"] {
        display: flex !important;
        justify-content: center !important;
    }

    /* Toggle switch background when OFF (not included) - RED */
    label[data-baseweb="checkbox"] div[role="switch"][aria-checked="false"] {
        background-color: #dc3545 !important;
    }

    /* Toggle switch background when ON (included) - GREEN */
    label[data-baseweb="checkbox"] div[role="switch"][aria-checked="true"] {
        background-color: #28a745 !important;
    }

    /* Hover effects */
    label[data-baseweb="checkbox"] div[role="switch"]:hover {
        opacity: 0.9 !important;
    }
</style>
""", unsafe_allow_html=True)


def clean_html_content(text):
    """Clean HTML tags and decode HTML entities from text."""
    if not text:
        return ""

    # Decode HTML entities (like &#x27; to ')
    text = html.unescape(text)

    # Remove HTML tags but keep the content
    text = re.sub(r'<br\s*/?>', '\n', text)  # Convert <br> to newlines
    text = re.sub(r'<p>', '\n', text)  # Convert <p> to newlines
    text = re.sub(r'</p>', '\n', text)  # Convert </p> to newlines
    text = re.sub(r'<h[1-6]>', '\n### ', text)  # Convert headings
    text = re.sub(r'</h[1-6]>', '\n', text)  # Close headings
    text = re.sub(r'<a\s+href=["\']([^"\']+)["\'][^>]*>([^<]+)</a>', r'\2 (\1)', text)  # Convert links to text with URL
    text = re.sub(r'<[^>]+>', '', text)  # Remove all other HTML tags

    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
    text = text.strip()

    return text


@st.cache_resource
def get_db_session():
    """Get database session (cached)."""
    return get_session()


def load_articles(session, filters):
    """Load articles with filters."""
    query = session.query(Article)

    # Apply filters
    if filters['category'] != 'All':
        category = session.query(Category).filter_by(name=filters['category']).first()
        if category:
            query = query.filter_by(category_id=category.id)

    if filters['status'] != 'All':
        query = query.filter_by(status=filters['status'])

    if filters['industry'] != 'All':
        industry = session.query(Industry).filter_by(name=filters['industry']).first()
        if industry:
            query = query.filter(Article.industries.contains(industry))

    # Date range filter
    if filters.get('date_from'):
        query = query.filter(Article.published_date >= filters['date_from'])
    if filters.get('date_to'):
        query = query.filter(Article.published_date <= filters['date_to'])

    # Order by most recent first
    query = query.order_by(Article.published_date.desc())

    return query.all()


def load_newsletters(session, filters):
    """Load newsletters with filters."""
    query = session.query(Newsletter)

    # Apply filters
    if filters['status'] != 'All':
        query = query.filter_by(status=filters['status'])

    if filters.get('source') and filters['source'] != 'All':
        query = query.filter_by(source=filters['source'])

    # Date range filter
    if filters.get('date_from'):
        query = query.filter(Newsletter.published_date >= filters['date_from'])
    if filters.get('date_to'):
        query = query.filter(Newsletter.published_date <= filters['date_to'])

    # Order by most recent first
    query = query.order_by(Newsletter.published_date.desc())

    newsletters = query.all()

    # Tag filter (post-query since tags are stored as JSON)
    if filters.get('tag') and filters['tag'] != 'All':
        newsletters = [n for n in newsletters if n.tags and filters['tag'] in n.tags]
    
    # Industry filter (post-query since industries are stored as JSON)
    if filters.get('industry') and filters['industry'] != 'All':
        newsletters = [n for n in newsletters if hasattr(n, 'industries') and n.industries and filters['industry'] in n.industries]

    return newsletters


def update_article_status(session, article_id, new_status, analysis_text=None):
    """Update article status and analysis."""
    article = session.query(Article).get(article_id)
    if article:
        article.status = new_status
        if analysis_text is not None:
            article.your_analysis = analysis_text
        session.commit()
        return True
    return False


def update_newsletter_status(session, newsletter_id, new_status, analysis_text=None):
    """Update newsletter status and analysis."""
    newsletter = session.query(Newsletter).get(newsletter_id)
    if newsletter:
        newsletter.status = new_status
        if analysis_text is not None:
            newsletter.your_analysis = analysis_text
        session.commit()
        return True
    return False


def render_article_review_page(session):
    """Original article review page."""
    st.title("Article Review")
    st.markdown("Review and mark articles for inclusion in the newsletter.")

    # Sidebar - Filters and Stats
    st.sidebar.header("Filters")

    # Get filter options
    categories = ['All'] + [cat.name for cat in session.query(Category).all()]
    industries = ['All'] + [ind.name for ind in session.query(Industry).all()]
    statuses = ['All', 'included', 'not_included', 'in_review', 'generated', 'finalized']

    # Status display mapping
    status_display = {
        'All': 'All',
        'included': 'Included',
        'not_included': 'Not Included', 
        'in_review': 'In Review',
        'generated': 'Generated',
        'finalized': 'Finalized'
    }

    # Filter controls
    filters = {
        'category': st.sidebar.selectbox("Category", categories, key="filter_category"),
        'industry': st.sidebar.selectbox("Industry", industries, key="filter_industry"),
        'status': st.sidebar.selectbox("Status", statuses, format_func=lambda x: status_display[x], key="filter_status"),
    }

    # Date range filter
    with st.sidebar.expander("Date Range)"):
        date_from = st.date_input("From", value=None, key="date_from")
        date_to = st.date_input("To", value=None, key="date_to")
        if date_from:
            filters['date_from'] = datetime.combine(date_from, datetime.min.time())
        if date_to:
            filters['date_to'] = datetime.combine(date_to, datetime.max.time())

    # Apply Filters button
    if st.sidebar.button("Apply Filters", type="primary", use_container_width=True):
        st.session_state.applied_filters = filters.copy()
        st.session_state.filter_applied = True
        st.session_state.articles_to_show = 20  # Reset to 20 when applying new filters
        st.rerun()

    # Clear Filters button
    if st.sidebar.button("Clear Filters", use_container_width=True):
        st.session_state.applied_filters = {
            'category': 'All',
            'industry': 'All',
            'status': 'All'
        }
        st.session_state.filter_applied = True
        st.session_state.articles_to_show = 20  # Reset to 20 when clearing filters
        st.rerun()

    # Use applied filters or default to 'All' on first load
    if st.session_state.applied_filters is None:
        # First load - default to showing all
        active_filters = {
            'category': 'All',
            'industry': 'All',
            'status': 'All'
        }
        st.session_state.applied_filters = active_filters
    else:
        active_filters = st.session_state.applied_filters

    # Statistics
    st.sidebar.header("Statistics")
    total_articles = session.query(Article).count()
    
    # Count articles in workflow (included, in_review, generated, finalized)
    in_workflow_count = session.query(Article).filter(
        Article.status.in_(['included', 'in_review', 'generated', 'finalized'])
    ).count()
    
    # Count by specific status
    included_count = session.query(Article).filter_by(status='included').count()
    in_review_count = session.query(Article).filter_by(status='in_review').count()
    generated_count = session.query(Article).filter_by(status='generated').count()
    finalized_count = session.query(Article).filter_by(status='finalized').count()
    not_included_count = session.query(Article).filter_by(status='not_included').count()

    st.sidebar.metric("Total Articles", total_articles)
    st.sidebar.metric("In Workflow", in_workflow_count)
    st.sidebar.metric("Not Included", not_included_count)
    
    # Show breakdown in expander
    with st.sidebar.expander("Workflow Breakdown"):
        st.write(f"Included: {included_count}")
        st.write(f"In Review: {in_review_count}")
        st.write(f"Generated: {generated_count}")
        st.write(f"Finalized: {finalized_count}")

    # Main content
    all_articles = load_articles(session, active_filters)
    total_articles = len(all_articles)

    # Show active filters
    filter_info = []
    if active_filters['category'] != 'All':
        filter_info.append(f"Category: {active_filters['category']}")
    if active_filters['industry'] != 'All':
        filter_info.append(f"Industry: {active_filters['industry']}")
    if active_filters['status'] != 'All':
        filter_info.append(f"Status: {status_display[active_filters['status']]}")
    if active_filters.get('date_from') or active_filters.get('date_to'):
        filter_info.append("Date range applied")

    if filter_info:
        st.info(f"Active filters: {' | '.join(filter_info)}")

    st.header(f"Articles ({total_articles} found)")

    if not all_articles:
        st.info("No articles found with current filters.")
        return

    # Pagination - show limited articles
    articles_to_display = all_articles[:st.session_state.articles_to_show]
    showing_count = len(articles_to_display)

    st.write(f"Showing {showing_count} of {total_articles} articles")

    # Display articles with card design (matching newsletter style)
    for i, article in enumerate(articles_to_display):
        with st.container():
            # Prepare summary with properly rendered markdown links
            if article.summary:
                summary = article.summary
                # Convert markdown links to HTML for proper rendering
                import re
                summary_html = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', summary)
            else:
                summary_html = '<em>No summary available</em>'

            # Build tags HTML
            tags_html = ''
            if article.industries:
                # Map industries to tag colors (reuse newsletter tag classes)
                industry_class_map = {
                    'General AI': 'tools',
                    'Healthcare and Medicine': 'research',
                    'Finance and Banking': 'funding',
                    'Legal and LegalTech': 'legal',
                    'Cybersecurity': 'security',
                    'Software and Development': 'tools',
                    'Manufacturing and Robotics': 'tools',
                    'Marketing and Advertising': 'funding',
                    'Education and EdTech': 'research',
                    'Retail and E-commerce': 'funding',
                    'Transportation and Autonomous Vehicles': 'tools',
                    'Energy and Utilities': 'research',
                    'Media and Entertainment': 'tools',
                    'Agriculture and AgriTech': 'research',
                    'Real Estate and PropTech': 'funding',
                    'Customer Service and Support': 'tools',
                    'Human Resources and Recruiting': 'funding',
                    'Supply Chain and Logistics': 'tools',
                    'Telecommunications': 'tools',
                    'Government and Public Sector': 'legal',
                }

                tags = [(ind.name, industry_class_map.get(ind.name, 'tools')) for ind in article.industries[:3]]  # Show max 3
                tags_html = '<div class="tag-container">' + ''.join([f'<span class="tag {tag[1]}">{tag[0]}</span>' for tag in tags]) + '</div>'

            # Create two columns: content and button
            col_content, col_button = st.columns([0.90, 0.10])

            with col_content:
                # Title HTML with embedded link
                title_html = f'<a href="{article.url}" class="newsletter-title-link" target="_blank">{article.title}</a>'

                # Render card content
                card_html = f'''
                <div class="newsletter-card">
                    <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 12px; flex-wrap: wrap;">
                        {title_html}
                        {tags_html}
                    </div>
                    <div class="newsletter-preview-text">{summary_html}</div>
                '''
                st.markdown(card_html, unsafe_allow_html=True)

            with col_button:
                # Status indicator button - shows current state
                is_included = article.status == 'included'
                btn_color = "#28a745" if is_included else "#dc3545"
                btn_symbol = "✓" if is_included else "×"
                help_text = "Included - Click to exclude" if is_included else "Excluded - Click to include"
                new_status = 'not_included' if is_included else 'included'

                # Add inline CSS for button styling
                st.markdown(f'''
                <style>
                button[key="btn_toggle_article_{article.id}"] {{
                    background-color: {btn_color} !important;
                    background: {btn_color} !important;
                    color: white !important;
                    border: 2px solid {btn_color} !important;
                    border-radius: 8px !important;
                    width: 50px !important;
                    min-width: 50px !important;
                    height: 50px !important;
                    min-height: 50px !important;
                    font-size: 24px !important;
                    font-weight: bold !important;
                }}
                </style>
                ''', unsafe_allow_html=True)

                if st.button(btn_symbol, key=f"btn_toggle_article_{article.id}", help=help_text):
                    update_article_status(session, article.id, new_status, article.your_analysis)
                    st.rerun()

            # Continue card content in the same column context
            with col_content:
                # Expandable section for notes only
                with st.expander("Save Notes", expanded=False):
                    # User analysis
                    st.markdown("**Your Analysis:**")
                    analysis = st.text_area(
                        "Add your notes:",
                        value=article.your_analysis or "",
                        height=100,
                        key=f"analysis_{article.id}",
                        placeholder="Add your thoughts, insights, or notes...",
                        label_visibility="collapsed"
                    )

                    if st.button("Save Notes", key=f"save_{article.id}", type="primary", use_container_width=True):
                        update_article_status(session, article.id, article.status, analysis)
                        st.success("Notes saved!")
                        st.rerun()

                # Close card div
                st.markdown('</div>', unsafe_allow_html=True)

    # Load More button
    if showing_count < total_articles:
        remaining = total_articles - showing_count
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(f"Load More ({remaining} remaining)", type="primary", use_container_width=True):
                st.session_state.articles_to_show += 20  # Load 20 more
                st.rerun()


def format_newsletter_tags(newsletter):
    """
    Format newsletter tags for display.
    Reads tags from database and maps to (display, class) tuples.
    """
    if not newsletter.tags:
        return []

    # Map category names to CSS classes
    tag_class_map = {
        'Tools and Products': 'tools',
        'Legal and Regulations': 'legal',
        'Funding and M&A': 'funding',
        'Security and Privacy': 'security',
        'Events and Conferences': 'events',
        'Research and Studies': 'research'
    }

    # Convert to (display, class) tuples
    tags = [(tag, tag_class_map.get(tag, 'tools')) for tag in newsletter.tags]
    return tags


def render_newsletter_review_page(session):
    """Newsletter review page with modern card-based UI."""
    # Header
    st.markdown('<h1 style="font-size: 28px; font-weight: 600; margin-bottom: 8px;">Newsletters</h1>', unsafe_allow_html=True)

    # Sidebar - Filters and Stats
    st.sidebar.header("Filters")

    # Get newsletter sources
    sources = ['All'] + [row[0] for row in session.query(Newsletter.source).distinct().all()]
    statuses = ['All', 'included', 'not_included', 'in_review', 'generated', 'finalized']

    # Status display mapping
    status_display = {
        'All': 'All',
        'included': 'Included',
        'not_included': 'Not Included',
        'in_review': 'In Review',
        'generated': 'Generated',
        'finalized': 'Finalized'
    }

    # Get available tags and industries from config
    import config
    available_tags = ['All'] + config.CATEGORIES
    available_industries = ['All'] + config.INDUSTRIES

    # Filter controls - Order: Category Tag, Industry, Source, Date
    filters = {
        'tag': st.sidebar.selectbox("Category Tag", available_tags, key="filter_newsletter_tag"),
        'industry': st.sidebar.selectbox("Industry", available_industries, key="filter_newsletter_industry"),
        'source': st.sidebar.selectbox("Source", sources, key="filter_newsletter_source"),
        'status': st.sidebar.selectbox("Status", statuses, format_func=lambda x: status_display[x], key="filter_newsletter_status"),
    }

    # Date range filter
    with st.sidebar.expander("Date Rangea"):
        date_from = st.date_input("From", value=None, key="newsletter_date_from")
        date_to = st.date_input("To", value=None, key="newsletter_date_to")
        if date_from:
            filters['date_from'] = datetime.combine(date_from, datetime.min.time())
        if date_to:
            filters['date_to'] = datetime.combine(date_to, datetime.max.time())

    # Apply Filters button
    if st.sidebar.button("Apply Filters", type="primary", use_container_width=True, key="apply_newsletter_filters"):
        st.session_state.applied_filters = filters.copy()
        st.session_state.filter_applied = True
        st.session_state.newsletters_to_show = 50  # Reset pagination
        st.rerun()

    # Clear Filters button
    if st.sidebar.button("Clear Filters", use_container_width=True, key="clear_newsletter_filters"):
        st.session_state.applied_filters = {
            'tag': 'All',
            'industry': 'All',
            'source': 'All',
            'status': 'All'
        }
        st.session_state.filter_applied = True
        st.session_state.newsletters_to_show = 50  # Reset pagination
        st.rerun()

    # Use applied filters
    if st.session_state.applied_filters is None:
        active_filters = {'tag': 'All', 'industry': 'All', 'source': 'All', 'status': 'All'}
        st.session_state.applied_filters = active_filters
    else:
        active_filters = st.session_state.applied_filters

    # Statistics
    st.sidebar.header("Statistics")
    total_newsletters = session.query(Newsletter).count()
    included_count = session.query(Newsletter).filter_by(status='included').count()
    not_included_count = session.query(Newsletter).filter_by(status='not_included').count()

    st.sidebar.metric("Total Newsletters", total_newsletters)
    st.sidebar.metric("Included", included_count)
    st.sidebar.metric("Not Included", not_included_count)

    # Main content
    all_newsletters = load_newsletters(session, active_filters)
    total = len(all_newsletters)

    # Pagination - show limited newsletters
    newsletters_to_display = all_newsletters[:st.session_state.newsletters_to_show]
    showing_count = len(newsletters_to_display)

    # Count display
    st.markdown(f'<p style="color: #8b92a8; font-size: 16px; margin-bottom: 24px;">Showing {showing_count} of {total} newsletters</p>', unsafe_allow_html=True)

    if not all_newsletters:
        st.info("No newsletters found with current filters.")
        return

    # Display newsletters with modern card design
    for i, newsletter in enumerate(newsletters_to_display):
        # Use container and wrap with styled div
        with st.container():
            # Build complete card content as HTML for title, meta, summary, tags
            archive_link = newsletter.archive_url if hasattr(newsletter, 'archive_url') and newsletter.archive_url else None
            date_str = newsletter.published_date.strftime("%b %d, %Y") if newsletter.published_date else "N/A"
            status_badge_class = "status-badge-read" if newsletter.status == 'not_included' else "status-badge"
            status_text = "Included" if newsletter.status == 'included' else "Read"

            # Prepare summary with properly rendered markdown links (no truncation)
            if newsletter.summary:
                summary = newsletter.summary
                # Convert markdown links to HTML for proper rendering
                import re
                summary_html = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', summary)
            else:
                summary_html = '<em>No summary available</em>'

            # Tags HTML
            tags = format_newsletter_tags(newsletter)
            tags_html = ''
            if tags:
                tags_html = '<div class="tag-container">' + ''.join([f'<span class="tag {tag[1]}">{tag[0]}</span>' for tag in tags]) + '</div>'

            # Create two columns: content and button
            col_content, col_button = st.columns([0.90, 0.10])

            with col_content:
                # Title HTML - make it a link only if archive_url exists
                if archive_link:
                    title_html = f'<a href="{archive_link}" class="newsletter-title-link" target="_blank">{newsletter.title}</a>'
                else:
                    title_html = f'<span class="newsletter-title-link" style="cursor: default;">{newsletter.title}</span>'

                # Render all content as one HTML block inside card (tags next to title)
                card_html = f'''
                <div class="newsletter-card">
                    <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 12px; flex-wrap: wrap;">
                        {title_html}
                        {tags_html}
                    </div>
                    <div class="newsletter-preview-text">{summary_html}</div>
                '''
                st.markdown(card_html, unsafe_allow_html=True)

            with col_button:
                # Status indicator button - shows current state
                # Consider 'generated' as included for display purposes
                is_included = newsletter.status in ['included', 'generated', 'finalized']
                btn_color = "#28a745" if is_included else "#dc3545"
                btn_symbol = "✓" if is_included else "×"
                help_text = "Included - Click to exclude" if is_included else "Excluded - Click to include"
                new_status = 'not_included' if is_included else 'included'

                # Create wrapper div with data attribute for CSS targeting
                st.markdown(f'<div class="newsletter-toggle-btn" data-newsletter-id="{newsletter.id}" data-color="{btn_color}" data-symbol="{btn_symbol}"></div>', unsafe_allow_html=True)
                
                # Add inline CSS as fallback to ensure color applies
                st.markdown(f'''
                <style>
                /* Aggressive CSS to override Streamlit defaults */
                button[key="btn_toggle_{newsletter.id}"],
                button[data-testid][key="btn_toggle_{newsletter.id}"],
                button[data-testid="baseButton-secondary"],
                button[data-testid="baseButton-primary"] {{
                    background-color: {btn_color} !important;
                    background: {btn_color} !important;
                    background-image: none !important;
                    color: white !important;
                    border: 2px solid {btn_color} !important;
                    border-radius: 8px !important;
                    width: 50px !important;
                    min-width: 50px !important;
                    height: 50px !important;
                    min-height: 50px !important;
                    padding: 0 !important;
                    font-size: 24px !important;
                    font-weight: bold !important;
                }}
                /* Target by parent container */
                div[data-testid="column"]:nth-of-type(2) button {{
                    background-color: {btn_color} !important;
                    background: {btn_color} !important;
                }}
                </style>
                ''', unsafe_allow_html=True)
                
                if st.button(btn_symbol, key=f"btn_toggle_{newsletter.id}", help=help_text):
                    update_newsletter_status(session, newsletter.id, new_status)
                    st.rerun()

            # Continue card content in the same column context
            with col_content:
                # Expandable section for details
                with st.expander("View Details and Links", expanded=False):
                    # Show extracted links
                    if newsletter.extracted_links:
                        st.markdown("**Links in this newsletter:**")
                        st.markdown("")

                        for idx, link in enumerate(newsletter.extracted_links[:10], 1):  # Show first 10
                            link_title = link.get('title', 'Link')
                            link_url = link.get('url', '')
                            link_context = link.get('context', '')

                            st.markdown(f"**{idx}.** [{link_title}]({link_url})")

                            if link_context and len(link_context) > 10:
                                clean_context = link_context[:150] + "..." if len(link_context) > 150 else link_context
                                st.markdown(f"> {clean_context}")

                            st.markdown("")

                        if len(newsletter.extracted_links) > 10:
                            st.markdown(f"*... and {len(newsletter.extracted_links) - 10} more links*")

                        # View source article button (only if archive_link exists)
                        if archive_link:
                            st.markdown(f'<a href="{archive_link}" target="_blank" style="color: #6bb6ff; text-decoration: none; font-size: 14px;">View Original Newsletter</a>', unsafe_allow_html=True)
                    else:
                        st.info("No links extracted from this newsletter")

                    # Editor notes
                    st.markdown("---")
                    st.markdown("**Your Analysis:**")
                    analysis = st.text_area(
                        "Add your notes:",
                        value=newsletter.your_analysis or "",
                        height=100,
                        key=f"analysis_{newsletter.id}",
                        placeholder="Add your thoughts, insights, or notes...",
                        label_visibility="collapsed"
                    )

                    if st.button("Save Notes", key=f"save_{newsletter.id}", type="primary", use_container_width=True):
                        update_newsletter_status(session, newsletter.id, newsletter.status, analysis)
                        st.success("Notes saved!")
                        st.rerun()

                # Close card div
                st.markdown('</div>', unsafe_allow_html=True)
    
    # Load More button
    if showing_count < total:
        remaining = total - showing_count
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(f"Load More ({remaining} remaining)", type="primary", use_container_width=True, key="load_more_newsletters"):
                st.session_state.newsletters_to_show += 50  # Load 50 more
                st.rerun()


def render_review_and_generate_page(session):
    """Review page for generating final article."""
    st.title("Review & Generate")
    st.markdown("Review selected articles and generate your final newsletter article.")

    # Get all workflow articles (included, in_review, generated, finalized)
    included_articles = session.query(Article).filter(
        Article.status.in_(['included', 'in_review', 'generated', 'finalized'])
    ).order_by(
        Article.category_id, Article.published_date.desc()
    ).all()

    if not included_articles:
        st.warning("No articles in the workflow yet. Go to Article Review to include articles.")
        return

    # Display summary of included articles
    st.header(f"Included Articles ({len(included_articles)})")

    # Group by category
    articles_by_category = {}
    for article in included_articles:
        category_name = article.category.name if article.category else 'Uncategorized'
        if category_name not in articles_by_category:
            articles_by_category[category_name] = []
        articles_by_category[category_name].append(article)

    # Display articles grouped by category
    for category_name in sorted(articles_by_category.keys()):
        category_articles = articles_by_category[category_name]
        with st.expander(f"**{category_name}** ({len(category_articles)} articles)", expanded=True):
            for article in category_articles:
                # Show status badge
                status_display = article.status.upper().replace('_', ' ')
                st.markdown(f"### [{status_display}] {article.title}")
                st.markdown(f"**Source:** {article.source} | **Date:** {article.published_date.strftime('%Y-%m-%d') if article.published_date else 'N/A'}")

                if article.industries:
                    industries_str = ', '.join([ind.name for ind in article.industries])
                    st.markdown(f"**Industries:** {industries_str}")

                # Show summary
                if article.summary:
                    with st.expander("View Summary"):
                        st.write(article.summary)

                # Show user's analysis
                if article.your_analysis:
                    st.info(f"**Your Analysis:** {article.your_analysis}")

                st.markdown(f"[Read Original]({article.url})")
                st.divider()

    st.markdown("---")

    # User content/instructions section
    st.header("Add Your Content")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Additional Content")
        st.markdown("Add any additional context, data, or content you want to include:")
        user_content = st.text_area(
            "User Content",
            value=included_articles[0].user_content if included_articles[0].user_content else "",
            height=200,
            key="user_content_input",
            placeholder="Add your own insights, data, or additional context here..."
        )

    with col2:
        st.subheader("AI Generation Instructions")
        st.markdown("If you want AI to generate an article, provide instructions:")
        ai_instructions = st.text_area(
            "AI Instructions",
            value=included_articles[0].ai_instructions if included_articles[0].ai_instructions else "",
            height=200,
            key="ai_instructions_input",
            placeholder="E.g., 'Write a comprehensive article combining these insights, focusing on trends and implications...'"
        )

    st.markdown("---")

    # Action buttons
    st.header("Choose Your Path")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Option 1: Generate with AI")
        st.markdown("AI will create an article based on the selected articles, your analysis, and your instructions.")

        if st.button("Generate Article with AI", type="primary", use_container_width=True):
            if not ai_instructions.strip():
                st.error("Please provide AI instructions before generating.")
            else:
                # Save user content and instructions
                for article in included_articles:
                    article.user_content = user_content
                    article.ai_instructions = ai_instructions
                    article.status = 'in_review'
                session.commit()

                st.success("Instructions saved! Generating article...")

                # Call AI generation function
                generated_content = generate_article_with_ai(session, included_articles, user_content, ai_instructions)

                # Save generated content
                for article in included_articles:
                    article.generated_content = generated_content
                    article.status = 'generated'
                session.commit()

                st.success("Article generated! Go to 'Final Edit' to review and tweak.")
                st.session_state.current_page = 'final_edit'
                st.rerun()

    with col2:
        st.subheader("Option 2: Use as Final Report")
        st.markdown("Skip AI generation and use your current content as the final report.")

        if st.button("Use as Final Report", use_container_width=True):
            # Use standard format instead of compile_final_report
            final_content = generate_standard_format(included_articles, user_content)

            # Save as final content
            for article in included_articles:
                article.user_content = user_content
                article.final_content = final_content
                article.status = 'finalized'
            session.commit()

            st.success("Report finalized! Go to 'Final Edit' to view or make tweaks.")
            st.session_state.current_page = 'final_edit'
            st.rerun()


def render_final_edit_page(session):
    """Final editing page for tweaking generated/final content."""
    st.title("Final Edit")
    st.markdown("Review and edit your final article.")

    # Get articles that are generated or finalized
    articles = session.query(Article).filter(
        Article.status.in_(['generated', 'finalized'])
    ).all()

    if not articles:
        st.warning("No generated or finalized content yet. Go to 'Review & Generate' first.")
        return

    # Determine which content to show
    if articles[0].generated_content:
        current_content = articles[0].generated_content
        content_type = "AI Generated Article"
    elif articles[0].final_content:
        current_content = articles[0].final_content
        content_type = "Final Report"
    else:
        st.error("No content found.")
        return

    st.header(f"{content_type}")

    # Editable content
    edited_content = st.text_area(
        "Edit your content",
        value=current_content,
        height=500,
        key="final_content_editor"
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("Save Changes", type="primary", use_container_width=True):
            for article in articles:
                article.final_content = edited_content
                article.status = 'finalized'
            session.commit()
            st.success("Changes saved!")

    with col2:
        if st.button("Back to Review", use_container_width=True):
            st.session_state.current_page = 'review_generate'
            st.rerun()

    with col3:
        # Export format selector
        export_format = st.selectbox(
            "Export Format",
            ["Markdown (.md)", "Text (.txt)", "Word (.docx)", "PDF (.pdf)"],
            key="export_format_selector"
        )

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # Prepare content based on format
        if export_format == "Markdown (.md)":
            file_content = edited_content
            file_name = f"newsletter_article_{timestamp}.md"
            mime_type = "text/markdown"
        elif export_format == "Text (.txt)":
            # Strip markdown formatting for plain text
            import re
            text_content = re.sub(r'[#*_`\[\]()]', '', edited_content)
            file_content = text_content
            file_name = f"newsletter_article_{timestamp}.txt"
            mime_type = "text/plain"
        elif export_format == "Word (.docx)":
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from io import BytesIO

                doc = Document()

                # Process markdown into Word
                lines = edited_content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        # Heading 1
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        # Heading 2
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        # Heading 3
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        # Regular paragraph
                        doc.add_paragraph(line)

                # Save to BytesIO
                buffer = BytesIO()
                doc.save(buffer)
                file_content = buffer.getvalue()
                file_name = f"newsletter_article_{timestamp}.docx"
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except ImportError:
                st.error("python-docx library not installed. Install with: pip install python-docx")
                file_content = edited_content
                file_name = f"newsletter_article_{timestamp}.md"
                mime_type = "text/markdown"
        elif export_format == "PDF (.pdf)":
            try:
                from markdown import markdown
                from weasyprint import HTML
                from io import BytesIO

                # Convert markdown to HTML
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 40px; }}
                        h1 {{ color: #333; }}
                        h2 {{ color: #555; }}
                        h3 {{ color: #777; }}
                        p {{ line-height: 1.6; }}
                    </style>
                </head>
                <body>
                    {markdown(edited_content)}
                </body>
                </html>
                """

                # Convert HTML to PDF
                buffer = BytesIO()
                HTML(string=html_content).write_pdf(buffer)
                file_content = buffer.getvalue()
                file_name = f"newsletter_article_{timestamp}.pdf"
                mime_type = "application/pdf"
            except ImportError:
                st.error("PDF libraries not installed. Install with: pip install markdown weasyprint")
                file_content = edited_content
                file_name = f"newsletter_article_{timestamp}.md"
                mime_type = "text/markdown"

        # Download button
        st.download_button(
            label="Download",
            data=file_content,
            file_name=file_name,
            mime=mime_type,
            use_container_width=True
        )

    # Preview
    st.markdown("---")
    st.header("Preview")
    st.markdown(edited_content)


def compile_final_report(articles, user_content):
    """Compile articles and user content into a final report."""
    lines = []
    lines.append("# Newsletter Article\n")
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

    if user_content:
        lines.append("## Editor's Note\n\n")
        lines.append(f"{user_content}\n\n")

    # Group by category
    articles_by_category = {}
    for article in articles:
        category_name = article.category.name if article.category else 'Uncategorized'
        if category_name not in articles_by_category:
            articles_by_category[category_name] = []
        articles_by_category[category_name].append(article)

    # Add articles
    for category_name in sorted(articles_by_category.keys()):
        category_articles = articles_by_category[category_name]
        lines.append(f"## {category_name}\n\n")

        for article in category_articles:
            lines.append(f"### {article.title}\n\n")

            if article.your_analysis:
                lines.append(f"**Analysis:** {article.your_analysis}\n\n")

            if article.summary:
                lines.append(f"{article.summary}\n\n")

            lines.append(f"[Read more]({article.url})\n\n")

    return ''.join(lines)


def produce_all_links(newsletters):
    """
    Generate output with all links from all included newsletters.
    Each link gets a 1-line explanation with the link embedded.
    
    Args:
        newsletters: List of Newsletter objects
        
    Returns:
        str: Markdown formatted content with all links and explanations
    """
    lines = []
    lines.append("# Newsletter Links Summary\n")
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
    lines.append(f"**Total Newsletters:** {len(newsletters)}\n\n")
    lines.append("---\n\n")
    
    # Collect all links from all newsletters
    all_links = []
    for newsletter in newsletters:
        extracted_links = newsletter.extracted_links if newsletter.extracted_links else []
        all_links.extend(extracted_links)
    
    # Display all links in one unified list
    lines.append(f"**Total Links:** {len(all_links)}\n\n")
    
    for i, link in enumerate(all_links, 1):
        link_title = link.get('title', 'Untitled')
        link_url = link.get('url', '#')
        link_explanation = link.get('explanation', '')
        
        # Use explanation if available, otherwise use title
        if link_explanation and link_explanation != link_title:
            lines.append(f"{i}. {link_explanation} - [{link_title}]({link_url})\n")
        else:
            lines.append(f"{i}. [{link_title}]({link_url})\n")
    
    if not all_links:
        lines.append("*No links extracted from newsletters.*\n")
    
    return ''.join(lines)


def personalize_links_with_ai(newsletters, user_preferences, user_analysis=""):
    """
    Use AI to personalize and filter links based on user preferences and analysis.
    Generates 2-3 line paragraphs for each selected link.
    
    Args:
        newsletters: List of Newsletter objects
        user_preferences: str, User's description of what they find interesting
        user_analysis: str, Additional analysis/context from user
        
    Returns:
        str: AI-generated personalized content with selected links
    """
    # Check if OpenAI is available
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        st.error("OpenAI API key not found. Please set OPENAI_API_KEY environment variable.")
        return None
    
    try:
        client = openai.OpenAI(api_key=api_key)
        
        # Compile all newsletters and their links for context
        newsletter_context = []
        editor_notes = []
        
        for newsletter in newsletters:
            nl_data = {
                'title': newsletter.title,
                'source': newsletter.source,
                'summary': newsletter.summary,
                'your_analysis': newsletter.your_analysis,
                'links': newsletter.extracted_links if newsletter.extracted_links else []
            }
            newsletter_context.append(nl_data)
            
            # Collect editor notes
            if newsletter.your_analysis:
                editor_notes.append(f"- {newsletter.source}: {newsletter.your_analysis}")
        
        # Create context string
        context_lines = []
        
        # Add editor notes section
        if editor_notes:
            context_lines.append("=== EDITOR NOTES FROM REVIEW ===")
            context_lines.extend(editor_notes)
            context_lines.append("")
        
        # Add newsletter data
        context_lines.append("=== NEWSLETTERS ===")
        for i, nl in enumerate(newsletter_context, 1):
            context_lines.append(f"\n--- Newsletter {i}: {nl['source']} ---")
            context_lines.append(f"Title: {nl['title']}")
            if nl['summary']:
                context_lines.append(f"Summary: {nl['summary']}")
            if nl['your_analysis']:
                context_lines.append(f"Editor Notes: {nl['your_analysis']}")
            context_lines.append(f"\nLinks ({len(nl['links'])}):")
            for j, link in enumerate(nl['links'], 1):
                context_lines.append(f"{j}. {link.get('title', 'Untitled')} - {link.get('url', '#')}")
                if link.get('explanation'):
                    context_lines.append(f"   Explanation: {link.get('explanation')}")
        
        context_str = '\n'.join(context_lines)
        
        # Limit context length
        if len(context_str) > 10000:
            context_str = context_str[:10000] + "\n... (content truncated)"
        
        system_prompt = """You are an expert newsletter curator who personalizes content based on user interests and editor insights.

Your task:
1. Review all newsletters, editor notes, and links
2. Select ONLY the links that match the user's interests and analysis
3. For EACH selected link, write a 2-3 line paragraph explaining:
   - What the link is about
   - Why it's relevant to the user
   - Key insights or implications
4. Format each link as: [descriptive text](url) embedded naturally in the first sentence
5. Do NOT use emojis - plain text only

Format Requirements:
- List of links (numbered)
- Each link gets 2-3 lines of explanation
- Embed the link in the first sentence naturally
- Be concise but informative

Example format:
1. [OpenAI's new GPT-5 model](url) introduces breakthrough reasoning capabilities. This advancement directly addresses your interest in model architecture improvements. The implications for enterprise applications are significant.

2. [Google's latest computer vision research](url) demonstrates 40% improvement in accuracy. This aligns with your focus on CV breakthroughs. The technique is readily applicable to production systems."""

        # Build user prompt with all context
        user_prompt_parts = [f"User's Interests:\n{user_preferences}\n"]
        
        if user_analysis:
            user_prompt_parts.append(f"Additional Analysis/Context:\n{user_analysis}\n")
        
        user_prompt_parts.append(f"Newsletters Data:\n{context_str}\n")
        user_prompt_parts.append("Please curate and present only the most relevant links. For each link, provide 2-3 lines explaining what it is and why it matters based on the user's interests and the editor notes.")
        
        user_prompt = '\n'.join(user_prompt_parts)

        with st.spinner("🤖 AI is personalizing your links..."):
            response = client.chat.completions.create(
                model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=3000
            )
        
        personalized_content = response.choices[0].message.content.strip()
        
        # Add header
        header = f"# Personalized Newsletter Links\n\n"
        header += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        header += f"**Based on your interests:** {user_preferences}\n\n"
        if user_analysis:
            header += f"**Your Analysis:** {user_analysis}\n\n"
        header += "---\n\n"
        
        return header + personalized_content
        
    except Exception as e:
        st.error(f"Error generating personalized content: {e}")
        return None


def generate_report_for_selected_links(selected_links):
    """
    Generate a report with 2-3 sentences about each selected link using AI.
    
    Args:
        selected_links: List of link dictionaries with 'title', 'url', 'context'
        
    Returns:
        str: Formatted report with descriptions
    """
    # Check if OpenAI is available
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        st.error("OpenAI API key not found. Please set OPENAI_API_KEY environment variable.")
        return None
    
    try:
        client = openai.OpenAI(api_key=api_key)
        
        # Build context for AI
        links_context = []
        for i, link in enumerate(selected_links, 1):
            link_info = f"{i}. {link.get('title', 'Untitled')} - {link.get('url', '#')}"
            if link.get('context'):
                link_info += f"\n   Context: {link.get('context')[:200]}"
            links_context.append(link_info)
        
        context_str = '\n'.join(links_context)
        
        system_prompt = """You are an expert newsletter curator. For each link provided, write exactly 2-3 sentences that:
1. Explain what the link is about
2. Highlight the key information or value
3. Make it engaging and informative

Format:
- Start each entry with the number
- Write 2-3 complete sentences (no bullet points)
- Keep it concise and informative
- Use plain text only (no emojis)

Example format:
1. [Link Title](url) - This article discusses the latest breakthrough in AI reasoning capabilities. The research team achieved a 40% improvement in complex problem-solving tasks. This development could significantly impact enterprise AI applications.

2. [Another Link](url) - A comprehensive guide to implementing computer vision in production systems. The tutorial covers optimization techniques and real-world deployment strategies. Developers can apply these methods immediately to improve their CV pipelines."""

        user_prompt = f"""Please write 2-3 sentences for each of these links:

{context_str}

For each link, provide engaging and informative descriptions that explain what it's about and why it matters."""

        with st.spinner("🤖 Generating descriptions for selected links..."):
            response = client.chat.completions.create(
                model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
        
        generated_descriptions = response.choices[0].message.content.strip()
        
        # Add header
        header = f"# Newsletter Links Report\n\n"
        header += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        header += f"**Total Links:** {len(selected_links)}\n\n"
        header += "---\n\n"
        
        return header + generated_descriptions
        
    except Exception as e:
        st.error(f"Error generating report: {e}")
        return None


def is_unwanted_link(link):
    """
    Filter out sponsored, signup, and administrative links.
    
    Args:
        link: Dictionary with 'title', 'url', 'context'
        
    Returns:
        bool: True if link should be filtered out, False if it should be kept
    """
    title = link.get('title', '').lower()
    url = link.get('url', '').lower()
    
    # List of patterns to filter out
    unwanted_patterns = [
        # Sponsored/Ads
        'sponsor', 'sponsored', '(sponsor)', 'advertisement', 'advertise',
        
        # Sign up / Newsletter management
        'sign up', 'signup', 'subscribe', 'unsubscribe', 'manage your subscriptions',
        'manage subscriptions',
        
        # Legal/Admin
        'privacy policy', 'privacy', 'terms of use', 'terms', 'cookie policy',
        
        # Newsletter platform links
        'powered by', 'beehiiv', 'substack', 'mailchimp',
        
        # View/Share links
        'view online', 'view in browser', 'read online', 'web version',
        'share on twitter', 'share on facebook', 'share on linkedin',
        
        # Footer/Admin
        'manage preferences', 'update preferences', 'email preferences',
        'forward to a friend', 'forward this email',
        
        # Signature/Author links
        'thanks for reading', 'if you have any comments', 'just respond to this email',
        
        # Common author/editor names from TLDR (these are usually just signature links)
        'rachita kumar', 'matt cheung', 'gabriel sundaram',
        
        # Common marketing/promo
        'enterprise-grade', 'enterprising solutions',
    ]
    
    # ONLY filter truly unwanted admin/legal/promo patterns
    # Be less aggressive - only filter obvious spam/admin links
    for pattern in unwanted_patterns:
        if pattern in title or pattern in url:
            return True
    
    # Filter out platform domains
    unwanted_domains = [
        'beehiiv.com', 'substack.com', 'mailchimp.com',
        'facebook.com/sharer', 'twitter.com/intent', 'linkedin.com/sharing',
        'link.gettheelevator.com',  # Elevator tracking links
        'gettheelevator.com',  # Elevator links
    ]
    
    for domain in unwanted_domains:
        if domain in url:
            return True
    
    return False


def render_newsletter_review_and_generate_page(session):
    """Review page for generating final newsletter content from included newsletters."""
    st.title("Newsletter Review & Generate")
    st.markdown("Review included newsletters, select links, and generate your report.")
    
    # Get all included newsletters
    included_newsletters = session.query(Newsletter).filter(
        Newsletter.status.in_(['included', 'in_review', 'generated', 'finalized'])
    ).order_by(
        Newsletter.published_date.desc()
    ).all()
    
    if not included_newsletters:
        st.warning("No newsletters included yet. Go to Newsletter Review to include newsletters.")
        return
    
    # Display summary of included newsletters
    st.header(f"Included Newsletters ({len(included_newsletters)})")
    
    # Display newsletters summary in a compact way
    with st.expander(f"View {len(included_newsletters)} included newsletters", expanded=False):
        for newsletter in included_newsletters:
            st.markdown(f"**{newsletter.source}** - {newsletter.title[:100]}...")
            st.markdown(f"*Date:* {newsletter.published_date.strftime('%Y-%m-%d') if newsletter.published_date else 'N/A'} | *Links:* {len(newsletter.extracted_links) if newsletter.extracted_links else 0}")
            st.divider()
    
    st.markdown("---")
    
    # Collect all links from all newsletters and filter out unwanted ones
    all_links = []
    filtered_count = 0
    
    for newsletter in included_newsletters:
        if newsletter.extracted_links:
            for link in newsletter.extracted_links:
                # Check if link should be filtered out
                if is_unwanted_link(link):
                    filtered_count += 1
                    continue
                
                # Add source newsletter info to each link
                link_with_source = link.copy()
                link_with_source['source_newsletter'] = newsletter.source
                link_with_source['newsletter_title'] = newsletter.title
                all_links.append(link_with_source)
    
    if not all_links:
        st.warning("No content links found in the included newsletters (all links were filtered as spam/admin).")
        return
    
    # Show filtering info
    st.header(f"Content Links from Selected Newsletters ({len(all_links)} links)")
    if filtered_count > 0:
        st.info(f"✓ Automatically filtered out {filtered_count} administrative/sponsored links (privacy, terms, sign up, powered by, etc.)")
    st.markdown("**All links are selected by default.** Uncheck any links you want to exclude from the report.")
    
    # Initialize session state for link selection
    if 'selected_links' not in st.session_state:
        st.session_state.selected_links = {i: True for i in range(len(all_links))}
    
    # Make sure we have the right number of items in selected_links
    if len(st.session_state.selected_links) != len(all_links):
        st.session_state.selected_links = {i: True for i in range(len(all_links))}
    
    # Display all links with checkboxes
    st.markdown("---")
    
    # Add select/deselect all buttons
    col_all, col_none, col_count = st.columns([1, 1, 2])
    with col_all:
        if st.button("✓ Select All", use_container_width=True):
            st.session_state.selected_links = {i: True for i in range(len(all_links))}
            st.rerun()
    with col_none:
        if st.button("✗ Deselect All", use_container_width=True):
            st.session_state.selected_links = {i: False for i in range(len(all_links))}
            st.rerun()
    with col_count:
        selected_count = sum(1 for v in st.session_state.selected_links.values() if v)
        st.markdown(f"**Selected: {selected_count} / {len(all_links)}**")
    
    st.markdown("---")
    
    # Display links with checkboxes
    for i, link in enumerate(all_links):
        col_check, col_content = st.columns([0.05, 0.95])
        
        with col_check:
            # Checkbox for selection
            is_selected = st.checkbox(
                "",
                value=st.session_state.selected_links.get(i, True),
                key=f"link_checkbox_{i}",
                label_visibility="collapsed"
            )
            st.session_state.selected_links[i] = is_selected
        
        with col_content:
            link_title = link.get('title', 'Untitled Link')
            link_url = link.get('url', '#')
            link_context = link.get('context', '')
            source_nl = link.get('source_newsletter', 'Unknown')
            
            # Display link with source
            st.markdown(f"**{i+1}.** [{link_title}]({link_url})")
            st.markdown(f"*From: {source_nl}*")
            
            if link_context and len(link_context) > 10:
                clean_context = link_context[:200] + "..." if len(link_context) > 200 else link_context
                st.markdown(f"> {clean_context}")
            
            st.markdown("")
    
    st.markdown("---")
    
    # Generate button
    st.header("Generate Report")
    
    selected_count = sum(1 for v in st.session_state.selected_links.values() if v)
    
    if selected_count == 0:
        st.warning("Please select at least one link to generate a report.")
    else:
        st.markdown(f"Ready to generate a **StrictlyVC-style newsletter** with **{selected_count}** selected links.")
        st.markdown("The newsletter will be organized by category (Top News, Massive/Big/Smaller Fundings, Other News) with AI-generated descriptions.")
        
        if st.button("🤖 Generate Newsletter", type="primary", use_container_width=True):
            # Get selected links
            selected_links = [all_links[i] for i, selected in st.session_state.selected_links.items() if selected]

            if not selected_links:
                st.error("No links selected!")
                return

            # Generate StrictlyVC-style newsletter
            newsletter_title = f"AI & Tech Weekly - {datetime.now().strftime('%B %d, %Y')}"

            with st.spinner("🤖 Generating StrictlyVC-style newsletter..."):
                report_content = generate_strictlyvc_style_newsletter(selected_links, newsletter_title)

            if report_content:
                # Save generated content
                for newsletter in included_newsletters:
                    newsletter.generated_content = report_content
                    newsletter.status = 'generated'
                session.commit()

                st.success(f"✅ Newsletter generated with {len(selected_links)} links! Go to 'Newsletter Final Edit' to view.")
                st.session_state.current_page = 'newsletter_final_edit'
                st.rerun()


def render_newsletter_final_edit_page(session):
    """Final viewing page for newsletter content with export options."""
    st.title("Newsletter Final Report")
    st.markdown("Your generated report is ready! Export it as PDF or DOCX.")
    
    # Get newsletters that are generated or finalized
    newsletters = session.query(Newsletter).filter(
        Newsletter.status.in_(['generated', 'finalized'])
    ).all()
    
    if not newsletters:
        st.warning("No generated newsletter content yet. Go to 'Newsletter Review & Generate' first.")
        return
    
    # Get the generated content (should be the same for all in the batch)
    generated_content = newsletters[0].generated_content if newsletters[0].generated_content else ""
    final_content = newsletters[0].final_content if newsletters[0].final_content else generated_content
    
    # Mark as finalized if not already
    for newsletter in newsletters:
        if newsletter.status == 'generated':
            newsletter.final_content = final_content
            newsletter.status = 'finalized'
    session.commit()
    
    # Export buttons
    st.header("Export Options")
    
    col1, col2, col3 = st.columns(3)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    with col1:
        # Export as Markdown
        st.download_button(
            label="Export as Markdown",
            data=final_content,
            file_name=f"newsletter_report_{timestamp}.md",
            mime="text/markdown",
            use_container_width=True
        )
    
    with col2:
        # Export as DOCX
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            
            # Process markdown into Word
            lines = final_content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    # Handle markdown links in text
                    import re
                    # Simple handling - just keep the text for now
                    clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                    doc.add_paragraph(clean_line)
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            docx_data = buffer.getvalue()
            
            st.download_button(
                label="Export as Word",
                data=docx_data,
                file_name=f"newsletter_report_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except ImportError:
            st.button("Export as Word", disabled=True, use_container_width=True, help="Install python-docx: pip install python-docx")
    
    with col3:
        # Export as PDF - Use reportlab (simpler, works on Windows)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT
            from io import BytesIO
            import re
            
            # Create PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#333333',
                spaceAfter=12
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor='#555555',
                spaceAfter=8
            )
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=10,
                textColor='#000000',
                alignment=TA_LEFT,
                spaceAfter=10
            )
            
            # Build content
            story = []
            lines = final_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.1*inch))
                    continue
                
                if line.startswith('# '):
                    clean_line = line[2:].replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(clean_line, title_style))
                elif line.startswith('## '):
                    clean_line = line[3:].replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(clean_line, heading_style))
                elif line.startswith('### '):
                    clean_line = line[4:].replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(clean_line, heading_style))
                else:
                    # Handle markdown formatting
                    clean_line = line
                    
                    # First escape any existing HTML
                    clean_line = clean_line.replace('&', '&amp;')
                    
                    # Handle markdown links: [text](url)
                    clean_line = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2" color="blue">\1</a>', clean_line)
                    
                    # Handle bold: **text**
                    clean_line = re.sub(r'\*\*([^\*]+)\*\*', r'<b>\1</b>', clean_line)
                    
                    # Handle italic: *text*
                    clean_line = re.sub(r'\*([^\*]+)\*', r'<i>\1</i>', clean_line)
                    
                    # Escape any remaining < > that aren't part of our tags
                    # (This is already handled by the escaping above)
                    
                    try:
                        story.append(Paragraph(clean_line, body_style))
                    except Exception as e:
                        # If there's an error parsing, just add plain text
                        safe_text = line.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                        story.append(Paragraph(safe_text, body_style))
            
            # Build PDF
            doc.build(story)
            pdf_data = buffer.getvalue()
            
            st.download_button(
                label="Export as PDF",
                data=pdf_data,
                file_name=f"newsletter_report_{timestamp}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        except ImportError:
            st.button("Export as PDF", disabled=True, use_container_width=True, help="Install reportlab: pip install reportlab")
    
    # Preview section
    st.markdown("---")
    st.header("Report Preview")
    
    # Display as read-only
    st.markdown(final_content)


def main():
    """Main dashboard application."""
    st.sidebar.title("QS Newsletter Dashboard")

    session = get_db_session()

    # Workflow Mode Selector
    st.sidebar.markdown("---")
    st.sidebar.subheader("Workflow Mode")

    def on_content_type_change():
        """Callback when content type changes."""
        st.session_state.content_type = st.session_state.content_type_selector
        st.session_state.applied_filters = None  # Reset filters
        st.session_state.current_page = 'review'  # Go back to review page
        if st.session_state.content_type == 'newsletter_builder':
            st.session_state.builder_page = 'feed'  # Reset builder page

    content_type = st.sidebar.radio(
        "Select workflow:",
        options=['newsletters', 'articles', 'newsletter_builder'],
        format_func=lambda x: {
            'newsletter_builder': 'Newsletter Builder',
            'articles': 'Articles (Detailed)',
            'newsletters': 'Newsletters (Links)'
        }[x],
        key="content_type_selector",
        index={'newsletters': 0, 'articles': 1, 'newsletter_builder': 2}.get(
            st.session_state.content_type, 0
        ),
        on_change=on_content_type_change
    )

    # Update buttons for fetching new content (not shown in newsletter builder mode)
    if st.session_state.content_type != 'newsletter_builder':
        st.sidebar.markdown("---")
        st.sidebar.subheader("Update Content")

    if st.session_state.content_type == 'articles':
        if st.sidebar.button("Update Articles", type="primary", use_container_width=True):
            with st.spinner("Fetching articles from RSS feeds..."):
                try:
                    # Fetch articles from RSS feeds
                    articles = fetch_all_rss()

                    if articles:
                        # Save to temp file
                        save_fetched_articles(articles)

                        # Process and save to database
                        with st.spinner(f"Processing {len(articles)} articles..."):
                            processed_count = process_articles()

                        st.sidebar.success(f"✅ Added {processed_count} new articles!")
                        st.session_state.refresh_trigger += 1
                        st.rerun()
                    else:
                        st.sidebar.warning("No new articles found.")
                except Exception as e:
                    st.sidebar.error(f"Error updating articles: {e}")
    else:
        if st.sidebar.button("Update Newsletters", type="primary", use_container_width=True):
            with st.spinner("Fetching newsletters from Gmail..."):
                try:
                    # Fetch and process newsletters from Gmail (past 7 days)
                    processed_count = fetch_and_process_gmail_newsletters(days_back=7)

                    if processed_count > 0:
                        st.sidebar.success(f"✅ Added {processed_count} new newsletters!")
                        st.session_state.refresh_trigger += 1
                        st.rerun()
                    else:
                        st.sidebar.info("No new newsletters found.")
                except Exception as e:
                    st.sidebar.error(f"Error updating newsletters: {e}")

    # Clear All Inclusions button at top
    st.sidebar.markdown("---")
    if st.sidebar.button("Reset All Workflow", type="secondary", use_container_width=True):
        if st.session_state.content_type == 'articles':
            # Change all workflow articles back to "not_included"
            workflow_items = session.query(Article).filter(
                Article.status.in_(['included', 'in_review', 'generated', 'finalized'])
            ).all()
            for item in workflow_items:
                item.status = 'not_included'
            session.commit()
            st.sidebar.success(f"Reset {len(workflow_items)} articles to Not Included!")
        else:
            # Reset newsletters
            workflow_items = session.query(Newsletter).filter(
                Newsletter.status.in_(['included', 'in_review', 'generated', 'finalized'])
            ).all()
            for item in workflow_items:
                item.status = 'not_included'
            session.commit()
            st.sidebar.success(f"Reset {len(workflow_items)} newsletters to Not Included!")

        st.session_state.refresh_trigger += 1
        st.rerun()

    # Page navigation (only for articles and newsletters workflows, not newsletter builder)
    if st.session_state.content_type != 'newsletter_builder':
        st.sidebar.markdown("---")
        st.sidebar.subheader("Navigation")

        # Different pages for articles vs newsletters
        if st.session_state.content_type == 'articles':
            page_options = {
                'review': 'Article Review',
                'review_generate': 'Review & Generate',
                'final_edit': 'Final Edit'
            }
        else:
            page_options = {
                'review': 'Newsletter Review',
                'newsletter_review_generate': 'Newsletter Review & Generate',
                'newsletter_final_edit': 'Newsletter Final Edit'
            }

        # Use session state directly for the radio button value
        if 'current_page' not in st.session_state:
            st.session_state.current_page = 'review'

        def on_page_change():
            """Callback to update page when radio button changes."""
            st.session_state.current_page = st.session_state.page_selector

        selected_page = st.sidebar.radio(
            "Go to:",
            options=list(page_options.keys()),
            format_func=lambda x: page_options[x],
            key="page_selector",
            index=list(page_options.keys()).index(st.session_state.current_page) if st.session_state.current_page in page_options else 0,
            on_change=on_page_change
        )

    # Render selected page based on session state and content type
    if st.session_state.content_type == 'newsletter_builder':
        # NEW: Newsletter Builder workflow
        render_newsletter_builder_workflow(session)
    elif st.session_state.content_type == 'articles':
        if st.session_state.current_page == 'review':
            render_article_review_page(session)
        elif st.session_state.current_page == 'review_generate':
            render_review_and_generate_page(session)
        elif st.session_state.current_page == 'final_edit':
            render_final_edit_page(session)
    else:  # newsletters
        if st.session_state.current_page == 'review':
            render_newsletter_review_page(session)
        elif st.session_state.current_page == 'newsletter_preview':
            # Simplified preview page
            if st.session_state.preview_newsletter_id:
                render_newsletter_preview_page(session, st.session_state.preview_newsletter_id)
            else:
                st.error("No newsletter selected for preview")
                st.session_state.current_page = 'review'
                st.rerun()
        elif st.session_state.current_page == 'newsletter_review_generate':
            render_newsletter_review_and_generate_page(session)
        elif st.session_state.current_page == 'newsletter_final_edit':
            render_newsletter_final_edit_page(session)


if __name__ == "__main__":
    main()

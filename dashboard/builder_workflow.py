"""
Newsletter Builder Workflow
Unified content feed combining articles and newsletter links with manual control.
User has full control - AI is optional helper only.
"""

import streamlit as st
from datetime import datetime
from sqlalchemy.orm import Session

from src.utils.database import Article, Newsletter, Category, Industry


def get_unified_content(session: Session, filters: dict):
    """
    Fetch and combine articles + newsletter links into unified feed.
    Returns list of dicts with unified structure.
    """
    content_items = []
    
    # Fetch articles
    if filters.get('content_type') in ['all', 'articles']:
        article_query = session.query(Article)
        
        # Apply filters
        if filters.get('category') and filters['category'] != 'All':
            category = session.query(Category).filter_by(name=filters['category']).first()
            if category:
                article_query = article_query.filter_by(category_id=category.id)
        
        if filters.get('industry') and filters['industry'] != 'All':
            industry = session.query(Industry).filter_by(name=filters['industry']).first()
            if industry:
                article_query = article_query.filter(Article.industries.contains(industry))
        
        if filters.get('source') and filters['source'] != 'All':
            article_query = article_query.filter_by(source=filters['source'])
        
        if filters.get('date_from'):
            article_query = article_query.filter(Article.published_date >= filters['date_from'])
        if filters.get('date_to'):
            article_query = article_query.filter(Article.published_date <= filters['date_to'])
        
        articles = article_query.order_by(Article.published_date.desc()).limit(100).all()
        
        for article in articles:
            content_items.append({
                'id': f"article_{article.id}",
                'type': 'article',
                'db_id': article.id,
                'db_object': article,
                'title': article.title,
                'url': article.url,
                'source': article.source,
                'date': article.published_date,
                'summary': article.summary,
                'category': article.category.name if article.category else None,
                'industries': [ind.name for ind in article.industries] if article.industries else [],
                'commentary': getattr(article, 'commentary', None),
                'status': article.status,
                'newsletter_section': getattr(article, 'newsletter_section', None),
                'section_order': getattr(article, 'section_order', None)
            })
    
    # Fetch newsletters (extract individual links)
    if filters.get('content_type') in ['all', 'links']:
        newsletter_query = session.query(Newsletter)
        
        if filters.get('industry') and filters['industry'] != 'All':
            # Filter by industry (stored as JSON)
            newsletters = newsletter_query.all()
            newsletters = [n for n in newsletters if n.industries and filters['industry'] in n.industries]
        else:
            newsletters = newsletter_query.order_by(Newsletter.published_date.desc()).limit(50).all()
        
        # Extract links from newsletters
        link_counter = 0  # Add counter to ensure unique IDs
        for newsletter in newsletters:
            if newsletter.extracted_links:
                for link in newsletter.extracted_links:
                    # Skip if filtered by source
                    if filters.get('source') and filters['source'] != 'All' and newsletter.source != filters['source']:
                        continue
                    
                    link_counter += 1
                    content_items.append({
                        'id': f"link_{newsletter.id}_{link_counter}",  # Use counter for unique ID
                        'type': 'link',
                        'db_id': newsletter.id,
                        'db_object': newsletter,
                        'title': link.get('title', 'Untitled'),
                        'url': link.get('url', '#'),
                        'source': f"{newsletter.source} (Newsletter)",
                        'date': newsletter.published_date,
                        'summary': link.get('context', ''),
                        'category': None,
                        'industries': newsletter.industries if newsletter.industries else [],
                        'commentary': None,  # Links don't have individual commentary
                        'status': newsletter.status,
                        'newsletter_section': None,
                        'section_order': None
                    })
    
    # Sort by date (most recent first)
    content_items.sort(key=lambda x: x['date'] if x['date'] else datetime.min, reverse=True)
    
    return content_items


def render_content_feed(session: Session):
    """
    Main Content Feed page - unified view of articles and links.
    Simple, manual control - NO auto-AI.
    """
    st.title("Content Feed")
    st.markdown("**Manual curation mode** - Select content for your newsletter")
    
    # Filters in sidebar
    st.sidebar.header("Filters")
    
    # Content type filter
    content_type = st.sidebar.radio(
        "Show:",
        options=['all', 'articles', 'links'],
        format_func=lambda x: {'all': 'All Content', 'articles': 'Articles Only', 'links': 'Newsletter Links Only'}[x],
        key="builder_content_type"
    )
    
    # Get available sources
    article_sources = [row[0] for row in session.query(Article.source).distinct().all()]
    newsletter_sources = [row[0] for row in session.query(Newsletter.source).distinct().all()]
    all_sources = ['All'] + sorted(set(article_sources + newsletter_sources))
    
    filters = {
        'content_type': content_type,
        'source': st.sidebar.selectbox("Source", all_sources, key="builder_source"),
    }
    
    # Category/Industry filters (for articles)
    if content_type in ['all', 'articles']:
        categories = ['All'] + [cat.name for cat in session.query(Category).all()]
        filters['category'] = st.sidebar.selectbox("Category", categories, key="builder_category")
    
    # Industry filter (works for both)
    import config
    industries = ['All'] + config.INDUSTRIES
    filters['industry'] = st.sidebar.selectbox("Industry", industries, key="builder_industry")
    
    # Date range
    with st.sidebar.expander("Date Range"):
        date_from = st.date_input("From", value=None, key="builder_date_from")
        date_to = st.date_input("To", value=None, key="builder_date_to")
        if date_from:
            filters['date_from'] = datetime.combine(date_from, datetime.min.time())
        if date_to:
            filters['date_to'] = datetime.combine(date_to, datetime.max.time())
    
    if st.sidebar.button("Apply Filters", type="primary", use_container_width=True):
        st.session_state.builder_items_to_show = 50  # Reset pagination when filters change
        st.rerun()
    
    # Load content
    content_items = get_unified_content(session, filters)
    
    # Initialize session state for selected items
    if 'builder_selected_items' not in st.session_state:
        st.session_state.builder_selected_items = set()
    
    # Stats
    st.sidebar.markdown("---")
    st.sidebar.subheader("Statistics")
    articles_count = len([c for c in content_items if c['type'] == 'article'])
    links_count = len([c for c in content_items if c['type'] == 'link'])
    included_count = len(st.session_state.builder_selected_items)
    st.sidebar.metric("Total Items", len(content_items))
    st.sidebar.metric("Articles", articles_count)
    st.sidebar.metric("Links", links_count)
    st.sidebar.metric("Selected Items", included_count, delta=None)
    
    # Build Newsletter button in sidebar
    if included_count > 0:
        st.sidebar.markdown("---")
        st.sidebar.markdown("### ✨ Ready to Build?")
        if st.sidebar.button(f"🚀 Build Newsletter ({included_count} items)", type="primary", use_container_width=True, key="sidebar_build_btn"):
            st.session_state.builder_page = 'build'
            st.rerun()
    
    # Main content area
    total_items = len(content_items)
    
    if not content_items:
        st.info("No content found with current filters.")
        return
    
    # Pagination
    items_to_display = content_items[:st.session_state.builder_items_to_show]
    showing_count = len(items_to_display)
    
    st.header(f"Found {total_items} items")
    st.write(f"Showing {showing_count} of {total_items}")
    
    st.markdown("**Select items to include in your newsletter. You can add your own commentary for each item.**")
    st.markdown("---")
    
    # Display content items
    for item in items_to_display:
        with st.container():
            col_check, col_content = st.columns([0.05, 0.95])
            
            with col_check:
                # Checkbox to include/exclude (use session state)
                is_included = item['id'] in st.session_state.builder_selected_items
                include = st.checkbox(
                    "",
                    value=is_included,
                    key=f"include_{item['id']}",
                    label_visibility="collapsed"
                )
                
                # Update session state if changed
                if include and not is_included:
                    st.session_state.builder_selected_items.add(item['id'])
                elif not include and is_included:
                    st.session_state.builder_selected_items.discard(item['id'])
            
            with col_content:
                # Type badge
                badge_color = "#2196F3" if item['type'] == 'article' else "#FF9800"
                badge_text = "Article" if item['type'] == 'article' else "Link"
                
                st.markdown(f'<span style="background-color: {badge_color}; color: white; padding: 4px 12px; border-radius: 12px; font-size: 11px; font-weight: 600;">{badge_text}</span>', unsafe_allow_html=True)
                
                # Title
                st.markdown(f"### [{item['title']}]({item['url']})")
                
                # Meta info
                date_str = item['date'].strftime("%b %d, %Y") if item['date'] else "No date"
                industries_str = ", ".join(item['industries'][:3]) if item['industries'] else "No industries"
                st.markdown(f"**{item['source']}** | {date_str} | {industries_str}")
                
                # Summary preview
                if item['summary']:
                    summary = item['summary'][:200] + "..." if len(item['summary']) > 200 else item['summary']
                    st.markdown(f"> {summary}")
                
                # Commentary field (MANUAL - user writes their own)
                if include:
                    commentary = st.text_area(
                        "Your Commentary (optional):",
                        value=item['commentary'] or "",
                        height=80,
                        key=f"commentary_{item['id']}",
                        placeholder="Add your insights, take, or notes about this item...",
                        help="This will appear in your newsletter. Write what you want readers to know."
                    )
                    
                    # Save commentary button
                    if st.button("Save Commentary", key=f"save_comm_{item['id']}", type="secondary"):
                        db_obj = item['db_object']
                        db_obj.commentary = commentary
                        session.commit()
                        st.success("✓ Commentary saved")
                        st.rerun()
                
                st.markdown("---")
    
    # Load More button
    if showing_count < len(content_items):
        remaining = len(content_items) - showing_count
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(f"Load More ({remaining} remaining)", type="secondary", use_container_width=True, key="load_more_builder"):
                st.session_state.builder_items_to_show += 50  # Load 50 more
                st.rerun()
    
    # Build Newsletter button at bottom (same as sidebar)
    if included_count > 0:
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(f"🚀 Build Newsletter ({included_count} items)", type="primary", use_container_width=True, key="bottom_build_btn"):
                st.session_state.builder_page = 'build'
                st.rerun()


def render_build_newsletter_page(session: Session):
    """
    Newsletter Builder page - organize content into sections and export.
    Full manual control with optional AI assistance.
    """
    st.title("Build Your Newsletter")
    
    # Get selected items from session state
    if 'builder_selected_items' not in st.session_state or not st.session_state.builder_selected_items:
        st.warning("No content selected. Go back to Content Feed to select items.")
        if st.button("← Back to Content Feed"):
            st.session_state.builder_page = 'feed'
            st.rerun()
        return
    
    # Load the full content for selected items
    included_items = []
    
    for item_id in st.session_state.builder_selected_items:
        if item_id.startswith('article_'):
            article_id = int(item_id.split('_')[1])
            article = session.query(Article).get(article_id)
            if article:
                included_items.append({
                    'type': 'article',
                    'id': item_id,
                    'title': article.title,
                    'url': article.url,
                    'source': article.source,
                    'summary': article.summary,
                    'commentary': getattr(article, 'commentary', None)
                })
        elif item_id.startswith('link_'):
            # Extract newsletter_id and link_counter from id
            parts = item_id.split('_')
            newsletter_id = int(parts[1])
            link_counter = int(parts[2])
            
            newsletter = session.query(Newsletter).get(newsletter_id)
            if newsletter and newsletter.extracted_links:
                # Find the specific link by counter
                current_counter = 0
                for link in newsletter.extracted_links:
                    current_counter += 1
                    if current_counter == link_counter:
                        included_items.append({
                            'type': 'link',
                            'id': item_id,
                            'title': link.get('title', 'Untitled'),
                            'url': link.get('url', '#'),
                            'source': f"{newsletter.source} (Newsletter)",
                            'summary': link.get('context', ''),
                            'commentary': None
                        })
                        break
    
    total_items = len(included_items)
    
    if total_items == 0:
        st.warning("Selected items could not be loaded. Go back to Content Feed.")
        if st.button("← Back to Content Feed"):
            st.session_state.builder_page = 'feed'
            st.rerun()
        return
    
    # Count by type
    articles_count = len([item for item in included_items if item['type'] == 'article'])
    links_count = len([item for item in included_items if item['type'] == 'link'])
    
    st.markdown(f"**{total_items} items ready** ({articles_count} articles, {links_count} links)")

    # Initialize active view in session state
    if 'builder_active_view' not in st.session_state:
        st.session_state.builder_active_view = 'organize'

    # Check if we should switch to preview (after generation)
    if st.session_state.get('builder_switch_to_preview', False):
        st.session_state.builder_active_view = 'preview'
        st.session_state.builder_switch_to_preview = False

    # View selector
    st.markdown("---")
    view_options = {
        'organize': '📝 Organize Newsletter',
        'preview': '👁️ Preview & Export'
    }

    selected_view = st.radio(
        "Select view:",
        options=list(view_options.keys()),
        format_func=lambda x: view_options[x],
        index=0 if st.session_state.builder_active_view == 'organize' else 1,
        horizontal=True,
        key='builder_view_selector'
    )

    st.session_state.builder_active_view = selected_view
    st.markdown("---")

    # Render the selected view
    if selected_view == 'organize':
        st.header("Organize Your Newsletter")

        # Newsletter Title
        newsletter_title = st.text_input(
            "Newsletter Title",
            value=f"Newsletter - {datetime.now().strftime('%B %d, %Y')}",
            help="Main title for your newsletter"
        )
        
        # Introduction
        st.subheader("Introduction (Optional)")
        intro_text = st.text_area(
            "Write your intro:",
            value="",
            height=150,
            placeholder="Welcome to this week's newsletter! Here are the top stories and links I've curated for you...",
            help="Your opening message to readers"
        )

        st.markdown("---")

        # Content organization
        st.subheader("Your Content")

        # Display all items
        st.markdown(f"### Your Selected Items ({total_items})")
        for i, item in enumerate(included_items, 1):
            item_type = "📄 Article" if item['type'] == 'article' else "🔗 Link"
            with st.expander(f"{i}. {item_type}: {item['title'][:70]}...", expanded=False):
                st.markdown(f"**Source:** {item['source']}")
                st.markdown(f"**URL:** {item['url']}")

                if item.get('summary'):
                    st.markdown(f"**Summary:** {item['summary'][:150]}...")

                # Display commentary if available
                if item.get('commentary'):
                    st.markdown(f"**Your take:** {item['commentary']}")

        st.markdown("---")

        # Conclusion
        st.subheader("Conclusion (Optional)")
        outro_text = st.text_area(
            "Write your outro:",
            value="",
            height=150,
            placeholder="That's all for this week. Let me know your thoughts on these stories...",
            help="Your closing message to readers"
        )

        # Generate button
        st.markdown("---")
        col1, col2 = st.columns(2)

        with col1:
            if st.button("Generate Newsletter →", type="primary", use_container_width=True):
                # Store in session state
                st.session_state.newsletter_content = {
                    'title': newsletter_title,
                    'intro': intro_text,
                    'items': included_items,
                    'outro': outro_text,
                    'generated_date': datetime.now()
                }
                # Set flag to switch to preview tab
                st.session_state.builder_switch_to_preview = True
                st.success("Newsletter generated! Switching to preview...")
                st.rerun()

        with col2:
            if st.button("← Back to Content Feed", use_container_width=True):
                st.session_state.builder_page = 'feed'
                st.rerun()

    elif selected_view == 'preview':
        st.header("Newsletter Preview & Export")

        if 'newsletter_content' not in st.session_state:
            st.info("Click 'Generate Newsletter' in the Organize tab first.")
        else:
            content = st.session_state.newsletter_content
            final_md = generate_newsletter_markdown(content)

            # Export buttons at the top
            st.subheader("Export Newsletter")
            st.markdown("**Choose export format:**")

            col1, col2, col3 = st.columns(3)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            with col1:
                # Markdown export
                st.download_button(
                    label="📄 Export as Markdown",
                    data=final_md,
                    file_name=f"newsletter_{timestamp}.md",
                    mime="text/markdown",
                    type="primary",
                    use_container_width=True
                )

            with col2:
                # Word export
                try:
                    from docx import Document
                    from io import BytesIO
                    import re

                    doc = Document()

                    lines = final_md.split('\n')
                    for line in lines:
                        if line.startswith('# '):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        elif line.strip():
                            clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                            doc.add_paragraph(clean_line)

                    buffer = BytesIO()
                    doc.save(buffer)
                    docx_data = buffer.getvalue()

                    st.download_button(
                        label="📝 Export as Word",
                        data=docx_data,
                        file_name=f"newsletter_{timestamp}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                except ImportError:
                    st.button("📝 Export as Word", disabled=True, type="primary", use_container_width=True, help="Install python-docx")

            with col3:
                # PDF export
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.enums import TA_LEFT
                    from io import BytesIO
                    import re

                    buffer = BytesIO()
                    doc_pdf = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)

                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
                    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceAfter=8)
                    body_style = ParagraphStyle('CustomBody', parent=styles['BodyText'], fontSize=10, spaceAfter=10)

                    story = []
                    lines = final_md.split('\n')

                    for line in lines:
                        line = line.strip()
                        if not line:
                            story.append(Spacer(1, 0.1*inch))
                            continue

                        try:
                            if line.startswith('# '):
                                story.append(Paragraph(line[2:], title_style))
                            elif line.startswith('## '):
                                story.append(Paragraph(line[3:], heading_style))
                            elif line.startswith('### '):
                                story.append(Paragraph(line[4:], heading_style))
                            else:
                                clean_line = line.replace('&', '&amp;')
                                clean_line = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', clean_line)
                                clean_line = re.sub(r'\*\*([^\*]+)\*\*', r'<b>\1</b>', clean_line)
                                story.append(Paragraph(clean_line, body_style))
                        except:
                            pass

                    doc_pdf.build(story)
                    pdf_data = buffer.getvalue()

                    st.download_button(
                        label="📑 Export as PDF",
                        data=pdf_data,
                        file_name=f"newsletter_{timestamp}.pdf",
                        mime="application/pdf",
                        type="primary",
                        use_container_width=True
                    )
                except ImportError:
                    st.button("📑 Export as PDF", disabled=True, type="primary", use_container_width=True, help="Install reportlab")

            # Preview section
            st.markdown("---")
            st.subheader("Preview")
            st.markdown(final_md)


def generate_newsletter_markdown(content: dict) -> str:
    """
    Generate markdown newsletter from organized content.
    Purely from user's manual input - no AI generation.
    """
    lines = []
    
    # Title
    lines.append(f"# {content['title']}\n")
    lines.append(f"**Date:** {content['generated_date'].strftime('%B %d, %Y')}\n\n")
    lines.append("---\n\n")
    
    # Intro
    if content.get('intro'):
        lines.append(f"{content['intro']}\n\n")
        lines.append("---\n\n")
    
    # Content items
    if content.get('items'):
        lines.append("## Content\n\n")
        for i, item in enumerate(content['items'], 1):
            lines.append(f"### {i}. [{item['title']}]({item['url']})\n\n")
            lines.append(f"**Source:** {item['source']}\n\n")
            
            # User's commentary
            if item.get('commentary'):
                lines.append(f"{item['commentary']}\n\n")
            
            # Summary if no commentary
            elif item.get('summary'):
                summary = item['summary'][:300] + "..." if len(item['summary']) > 300 else item['summary']
                lines.append(f"{summary}\n\n")
            
            lines.append("\n")
    
    # Outro
    if content.get('outro'):
        lines.append("---\n\n")
        lines.append(f"{content['outro']}\n\n")
    
    return ''.join(lines)


def render_newsletter_builder_workflow(session: Session):
    """
    Main router for Newsletter Builder workflow.
    """
    # Initialize page state
    if 'builder_page' not in st.session_state:
        st.session_state.builder_page = 'feed'
    
    # Route to appropriate page
    if st.session_state.builder_page == 'feed':
        render_content_feed(session)
    elif st.session_state.builder_page == 'build':
        render_build_newsletter_page(session)

